from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length, Pt

from .models import Severity, ValidationIssue, ValidationReport

# --- Требования (можешь вынести в Settings/конфиг) ---
REQUIRED_FONT = "Times New Roman"
REQUIRED_SIZE_PT = 14

CAPTION_SIZE_PT = 12

# Поля 20 мм => в python-docx удобнее проверять в "twips" (1 inch = 1440 twips)
# 20 мм = 0.787401... inch => 1134 twips (примерно)
MM_TO_INCH = 1 / 25.4
TWIPS_PER_INCH = 1440
MARGIN_20MM_TWIPS = int(round(20 * MM_TO_INCH * TWIPS_PER_INCH))

# Допуски (Word иногда хранит значения не идеально)
MARGIN_TWIPS_TOL = 40  # ~0.7 мм
FONT_SIZE_TOL = 0.5    # pt
INDENT_CM = 1.25       # абзацный отступ
LINE_SPACING_TARGET = 1.15
LINE_SPACING_TOL = 0.06

# Объем 1-2 страницы. В DOCX нет "page count" без Word/LibreOffice,
# поэтому делаем грубую эвристику по количеству символов/абзацев.
MIN_TEXT_CHARS = 800     # примерно 1 стр с 14pt
MAX_TEXT_CHARS = 6500    # примерно 2 стр с 14pt


@dataclass(frozen=True)
class ThesisStyleRules:
    font_name: str = REQUIRED_FONT
    font_size_pt: float = REQUIRED_SIZE_PT
    margin_twips: int = MARGIN_20MM_TWIPS
    require_bold_header: bool = True
    title_uppercase: bool = True
    authors_italic: bool = True
    header_centered: bool = True
    body_justify: bool = True
    line_spacing: float = LINE_SPACING_TARGET
    paragraph_indent_cm: float = INDENT_CM
    require_literature_block: bool = True


def validate_thesis_docx(path: str | Path, rules: ThesisStyleRules | None = None) -> ValidationReport:
    rules = rules or ThesisStyleRules()
    doc = Document(str(path))

    issues: list[ValidationIssue] = []

    # 1) Параметры страницы: A4 + поля 20 мм, запрет колонтитулов/нумерации :contentReference[oaicite:2]{index=2}
    issues += _check_page_setup(doc, rules)

    # 2) Структура и заголовочная часть: title/authors/org по центру, title CAPS, authors italic :contentReference[oaicite:3]{index=3}
    issues += _check_header_structure(doc, rules)

    # 3) Основной текст: выравнивание по ширине, интервал 1.15, отступ 1.25 :contentReference[oaicite:4]{index=4}
    issues += _check_body_paragraphs(doc, rules)

    # 4) Литература: блок «Література» в конце, нумерация арабскими цифрами :contentReference[oaicite:5]{index=5}
    issues += _check_literature(doc, rules)

    # 5) Подписи рисунков/таблиц (по тексту): "Рис." / "Таблиця", 12 pt :contentReference[oaicite:6]{index=6}
    issues += _check_captions(doc)

    # 6) Объём 1-2 страницы (эвристика) :contentReference[oaicite:7]{index=7}
    issues += _check_length_heuristic(doc)

    ok = not any(i.severity == Severity.ERROR for i in issues)
    return ValidationReport(ok=ok, issues=issues)


# ---------------- internal checks ----------------

def _check_page_setup(doc: Document, rules: ThesisStyleRules) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    for idx, section in enumerate(doc.sections):
        # A4: 210x297 мм => 8.27 x 11.69 inch.
        # Точное сравнение сложно (twips), проверяем по диапазону.
        w = int(section.page_width)
        h = int(section.page_height)
        # A4 in twips: 8.27*1440=11909, 11.69*1440=16834
        if not (11700 <= w <= 12150 and 16600 <= h <= 17050):
            issues.append(ValidationIssue(
                code="PAGE_NOT_A4",
                severity=Severity.WARNING,
                message="Формат страницы отличается от A4 (ожидается 297x210 мм).",
                details={"section": idx, "page_width_twips": w, "page_height_twips": h},
            ))

        # Поля 20 мм все одинаковые :contentReference[oaicite:8]{index=8}
        for side_name, value in [
            ("top", int(section.top_margin)),
            ("bottom", int(section.bottom_margin)),
            ("left", int(section.left_margin)),
            ("right", int(section.right_margin)),
        ]:
            if abs(value - rules.margin_twips) > MARGIN_TWIPS_TOL:
                issues.append(ValidationIssue(
                    code="MARGIN_NOT_20MM",
                    severity=Severity.ERROR,
                    message="Поля страницы должны быть 20 мм со всех сторон.",
                    details={"section": idx, "side": side_name, "twips": value, "expected_twips": rules.margin_twips},
                ))

        # Запрет колонтитулов (проверяем, что в header/footer нет текста) :contentReference[oaicite:9]{index=9}
        if _section_has_header_footer_text(section):
            issues.append(ValidationIssue(
                code="HEADER_FOOTER_PRESENT",
                severity=Severity.ERROR,
                message="Колонтитулы запрещены (header/footer должен быть пустым).",
                details={"section": idx},
            ))

    return issues


def _section_has_header_footer_text(section) -> bool:
    def has_text(paragraphs: Iterable) -> bool:
        for p in paragraphs:
            if (p.text or "").strip():
                return True
        return False

    # header/footer могут быть None? в docx обычно есть объекты
    hdr = section.header
    ftr = section.footer
    if hdr and has_text(hdr.paragraphs):
        return True
    if ftr and has_text(ftr.paragraphs):
        return True
    return False


def _check_header_structure(doc: Document, rules: ThesisStyleRules) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []

    # Берем первые 3 непустых абзаца: title, authors, organization
    paragraphs = [p for p in doc.paragraphs if (p.text or "").strip()]
    if len(paragraphs) < 3:
        return [ValidationIssue(
            code="HEADER_TOO_SHORT",
            severity=Severity.ERROR,
            message="Не найдено минимум 3 строки заголовочной части: назва, автори, організація.",
        )]

    title_p, authors_p, org_p = paragraphs[0], paragraphs[1], paragraphs[2]

    # Центрирование заголовка/авторов/организации :contentReference[oaicite:10]{index=10}
    if rules.header_centered:
        for code, p, label in [
            ("TITLE_NOT_CENTER", title_p, "назва"),
            ("AUTHORS_NOT_CENTER", authors_p, "автори"),
            ("ORG_NOT_CENTER", org_p, "організація"),
        ]:
            if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                issues.append(ValidationIssue(
                    code=code,
                    severity=Severity.ERROR,
                    message=f"Строка «{label}» должна быть выровнена по центру.",
                    details={"text": (p.text or "").strip()},
                ))

    # Title uppercase :contentReference[oaicite:11]{index=11}
    title_text = (title_p.text or "").strip()
    if rules.title_uppercase and title_text and title_text != title_text.upper():
        issues.append(ValidationIssue(
            code="TITLE_NOT_UPPERCASE",
            severity=Severity.ERROR,
            message="Название должно быть оформлено прописными буквами (UPPERCASE).",
            details={"title": title_text},
        ))

    # Шрифт/кегль/жирность для заголовочной части :contentReference[oaicite:12]{index=12}
    issues += _check_paragraph_font(title_p, rules, must_bold=True, must_italic=False, where="title")
    issues += _check_paragraph_font(authors_p, rules, must_bold=True, must_italic=rules.authors_italic, where="authors")
    issues += _check_paragraph_font(org_p, rules, must_bold=True, must_italic=False, where="organization")

    # Авторы перечисляются через запятую, инициалы после фамилии (простая эвристика) :contentReference[oaicite:13]{index=13}
    if "," not in (authors_p.text or ""):
        issues.append(ValidationIssue(
            code="AUTHORS_LIST_FORMAT",
            severity=Severity.WARNING,
            message="Авторы обычно перечисляются через запятую.",
            details={"authors_line": (authors_p.text or "").strip()},
        ))
    # Эвристика "Фамилия І. О." (кириллица/латиница)
    if not re.search(r"[A-Za-zА-ЯІЇЄҐа-яіїєґ’\-]+\s+[A-ZА-ЯІЇЄҐ]\.\s*[A-ZА-ЯІЇЄҐ]\.", authors_p.text or ""):
        issues.append(ValidationIssue(
            code="AUTHORS_INITIALS_PATTERN",
            severity=Severity.WARNING,
            message="Проверь формат Фамилия И.О. (инициалы после фамилии).",
            details={"authors_line": (authors_p.text or "").strip()},
        ))

    return issues


def _check_paragraph_font(p, rules: ThesisStyleRules, must_bold: bool, must_italic: bool, where: str) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []

    # В DOCX шрифт может быть задан на run'ах, а в paragraph style.
    # Мы проверяем runs; если run без name/size — считаем "неизвестно" (warning).
    text = (p.text or "").strip()
    if not text:
        return issues

    # Собираем наблюдения по run'ам
    any_run = False
    bad_name = False
    bad_size = False
    bad_bold = False
    bad_italic = False
    unknown_font = False

    for r in p.runs:
        if not (r.text or "").strip():
            continue
        any_run = True

        name = r.font.name
        size = r.font.size

        if name is None:
            unknown_font = True
        else:
            if name != rules.font_name:
                bad_name = True

        if size is None:
            unknown_font = True
        else:
            if abs(size.pt - rules.font_size_pt) > FONT_SIZE_TOL:
                bad_size = True

        if must_bold and r.bold is not True:
            bad_bold = True

        if must_italic and r.italic is not True:
            bad_italic = True

    if not any_run:
        unknown_font = True

    if bad_name:
        issues.append(ValidationIssue(
            code="FONT_NAME_MISMATCH",
            severity=Severity.ERROR,
            message=f"Неверный шрифт в блоке {where}. Требуется {rules.font_name}.",
            details={"where": where, "text": text},
        ))
    if bad_size:
        issues.append(ValidationIssue(
            code="FONT_SIZE_MISMATCH",
            severity=Severity.ERROR,
            message=f"Неверный размер шрифта в блоке {where}. Требуется {rules.font_size_pt} pt.",
            details={"where": where, "text": text},
        ))
    if must_bold and bad_bold:
        issues.append(ValidationIssue(
            code="HEADER_NOT_BOLD",
            severity=Severity.ERROR,
            message=f"Блок {where} должен быть полужирным.",
            details={"where": where, "text": text},
        ))
    if must_italic and bad_italic:
        issues.append(ValidationIssue(
            code="AUTHORS_NOT_ITALIC",
            severity=Severity.ERROR,
            message="Список авторов должен быть курсивом.",
            details={"where": where, "text": text},
        ))
    if unknown_font and not (bad_name or bad_size):
        issues.append(ValidationIssue(
            code="FONT_UNKNOWN",
            severity=Severity.WARNING,
            message=f"Не удалось полностью определить шрифт/размер для блока {where} (возможно задан стилем).",
            details={"where": where},
        ))

    return issues


def _check_body_paragraphs(doc: Document, rules: ThesisStyleRules) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []

    paragraphs = [p for p in doc.paragraphs if (p.text or "").strip()]
    if len(paragraphs) < 4:
        return issues

    # Считаем, что body начинается после 3й строки (title/authors/org)
    body = paragraphs[3:]

    # Найдем "Література" и ограничим body до него
    lit_idx = None
    for i, p in enumerate(body):
        if (p.text or "").strip().lower() == "література":
            lit_idx = i
            break
    if lit_idx is not None:
        body = body[:lit_idx]

    # Проверяем первые N абзацев тела (чтобы не гонять весь документ)
    sample = body[:30] if len(body) > 30 else body
    for p in sample:
        txt = (p.text or "").strip()
        if not txt:
            continue

        # Выравнивание по ширине :contentReference[oaicite:14]{index=14}
        if rules.body_justify and p.alignment not in (None, WD_ALIGN_PARAGRAPH.JUSTIFY):
            issues.append(ValidationIssue(
                code="BODY_NOT_JUSTIFIED",
                severity=Severity.WARNING,
                message="Основной текст должен быть выровнен по ширине.",
                details={"paragraph": txt[:80]},
            ))

        # Межстрочный интервал 1.15 :contentReference[oaicite:15]{index=15}
        ls = p.paragraph_format.line_spacing
        if ls is not None:
            # line_spacing может быть float или Length
            if isinstance(ls, (int, float)):
                if abs(float(ls) - rules.line_spacing) > LINE_SPACING_TOL:
                    issues.append(ValidationIssue(
                        code="LINE_SPACING_MISMATCH",
                        severity=Severity.WARNING,
                        message="Ожидается межстрочный интервал 1.15.",
                        details={"got": float(ls), "expected": rules.line_spacing},
                    ))
            elif isinstance(ls, Length):
                # если задано в points — не сравниваем точно
                issues.append(ValidationIssue(
                    code="LINE_SPACING_ABSOLUTE",
                    severity=Severity.WARNING,
                    message="Межстрочный интервал задан абсолютным значением, ожидается множитель 1.15.",
                ))

        # Абзацный отступ 1.25 см :contentReference[oaicite:16]{index=16}
        indent = p.paragraph_format.first_line_indent
        if indent is not None:
            # 1.25 см -> в twips: cm->inch->twips
            expected_twips = int(round(rules.paragraph_indent_cm * 10 * MM_TO_INCH * TWIPS_PER_INCH))
            if abs(int(indent) - expected_twips) > 80:  # допуск ~1.4мм
                issues.append(ValidationIssue(
                    code="INDENT_MISMATCH",
                    severity=Severity.WARNING,
                    message="Абзацный отступ должен быть около 1.25 см.",
                    details={"got_twips": int(indent), "expected_twips": expected_twips},
                ))

        # Шрифт/кегль для тела (мягко: warning, потому что стиль может быть на уровне style)
        for r in p.runs:
            if not (r.text or "").strip():
                continue
            if r.font.name and r.font.name != rules.font_name:
                issues.append(ValidationIssue(
                    code="BODY_FONT_NAME_MISMATCH",
                    severity=Severity.WARNING,
                    message=f"Основной текст должен быть {rules.font_name}.",
                    details={"paragraph": txt[:80], "font": r.font.name},
                ))
                break
            if r.font.size and abs(r.font.size.pt - rules.font_size_pt) > FONT_SIZE_TOL:
                issues.append(ValidationIssue(
                    code="BODY_FONT_SIZE_MISMATCH",
                    severity=Severity.WARNING,
                    message=f"Основной текст должен быть {rules.font_size_pt} pt.",
                    details={"paragraph": txt[:80], "size_pt": r.font.size.pt},
                ))
                break

    return issues


def _check_literature(doc: Document, rules: ThesisStyleRules) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    paragraphs = [p for p in doc.paragraphs if (p.text or "").strip()]
    if not paragraphs:
        return issues

    # Ищем точный заголовок "Література" :contentReference[oaicite:17]{index=17}
    lit_pos = None
    for i, p in enumerate(paragraphs):
        if (p.text or "").strip().lower() == "література":
            lit_pos = i
            break

    if rules.require_literature_block and lit_pos is None:
        issues.append(ValidationIssue(
            code="LITERATURE_MISSING",
            severity=Severity.ERROR,
            message="Отсутствует блок «Література» в конце документа.",
        ))
        return issues

    if lit_pos is None:
        return issues

    # Литература должна быть ближе к концу
    if lit_pos < int(len(paragraphs) * 0.6):
        issues.append(ValidationIssue(
            code="LITERATURE_NOT_AT_END",
            severity=Severity.WARNING,
            message="Блок «Література» должен располагаться ближе к концу документа.",
            details={"position": lit_pos, "total_paragraphs": len(paragraphs)},
        ))

    # Проверим, что после заголовка есть хотя бы 1-2 пункта
    after = paragraphs[lit_pos + 1 : lit_pos + 6]
    has_item = any(re.match(r"^\s*\d+\.", (p.text or "").strip()) for p in after)
    if not has_item:
        issues.append(ValidationIssue(
            code="LITERATURE_ITEMS_FORMAT",
            severity=Severity.WARNING,
            message="Проверь оформление списка литературы: пункты обычно начинаются с «1.» и далее.",
        ))

    return issues


def _check_captions(doc: Document) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue

        # Рис. <номер> ... / Таблиця <номер> ...
        if re.match(r"^(Рис\.\s*\d+)", txt) or re.match(r"^(Таблиця\s*\d+)", txt):
            # Проверяем 12 pt на run'ах (если видно)
            for r in p.runs:
                if not (r.text or "").strip():
                    continue
                if r.font.size and abs(r.font.size.pt - CAPTION_SIZE_PT) > 0.5:
                    issues.append(ValidationIssue(
                        code="CAPTION_SIZE_MISMATCH",
                        severity=Severity.WARNING,
                        message="Подписи к рисункам/таблицам желательно оформлять 12 pt.",
                        details={"caption": txt[:80], "size_pt": r.font.size.pt},
                    ))
                    break
    return issues


def _check_length_heuristic(doc: Document) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    text = "\n".join((p.text or "") for p in doc.paragraphs).strip()
    chars = len(re.sub(r"\s+", " ", text))

    # Требование 1-2 страницы :contentReference[oaicite:18]{index=18}
    if chars < MIN_TEXT_CHARS:
        issues.append(ValidationIssue(
            code="LENGTH_TOO_SHORT",
            severity=Severity.WARNING,
            message="Документ выглядит слишком коротким для 1 полной страницы (эвристическая оценка).",
            details={"chars": chars, "min_expected": MIN_TEXT_CHARS},
        ))
    if chars > MAX_TEXT_CHARS:
        issues.append(ValidationIssue(
            code="LENGTH_TOO_LONG",
            severity=Severity.WARNING,
            message="Документ выглядит слишком длинным для 2 страниц (эвристическая оценка).",
            details={"chars": chars, "max_expected": MAX_TEXT_CHARS},
        ))
    return issues